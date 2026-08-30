from __future__ import annotations

from decimal import Decimal
from io import BytesIO
import re

from docx import Document
from docx.document import Document as DocumentType

from due_diligence_agent.domain.artifacts.models import Artifact, SourceLocator
from due_diligence_agent.domain.documents.models import (
    ParsedDocument,
    ParsedTable,
    ParserStatus,
    TextBlock,
)
from due_diligence_agent.ports.repositories import ArtifactStore


class DocxDocumentParser:
    parser_name = "python-docx"
    parser_version = "1"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(self, artifact_store: ArtifactStore) -> None:
        self._artifact_store = artifact_store

    def parse(self, artifact: Artifact, payload: bytes) -> ParsedDocument:
        try:
            document = Document(BytesIO(payload))
            paragraphs = self._paragraphs(artifact, document)
            tables = self._tables(artifact, document)
        except Exception:
            return ParsedDocument.outcome(
                artifact_id=artifact.id,
                detected_mime_type=self.media_type,
                parser_name=self.parser_name,
                parser_version=self.parser_version,
                status="damaged",
                error_code="damaged_docx",
            )
        all_blocks = paragraphs + [block for table in tables for block in table.text_blocks]
        status: ParserStatus = "parsed" if all_blocks else "partial"
        return ParsedDocument(
            artifact_id=artifact.id,
            detected_mime_type=self.media_type,
            tables=tables,
            text_blocks=all_blocks,
            parser_name=self.parser_name,
            parser_version=self.parser_version,
            confidence=Decimal("1") if status == "parsed" else Decimal("0.5"),
            status=status,
            error_code=None if status == "parsed" else "no_extractable_content",
        )

    def _paragraphs(self, artifact: Artifact, document: DocumentType) -> list[TextBlock]:
        blocks: list[TextBlock] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = _normalize(paragraph.text)
            if text:
                blocks.append(
                    self._store_text(
                        artifact,
                        text,
                        SourceLocator(
                            kind="docx_paragraph",
                            value=f"paragraph:{index}",
                            artifact_id=artifact.id,
                        ),
                    )
                )
        return blocks

    def _tables(self, artifact: Artifact, document: DocumentType) -> list[ParsedTable]:
        parsed: list[ParsedTable] = []
        for table_index, table in enumerate(document.tables, start=1):
            rows: list[list[str]] = []
            cells: list[TextBlock] = []
            for row_index, row in enumerate(table.rows, start=1):
                values: list[str] = []
                for cell_index, cell in enumerate(row.cells, start=1):
                    text = _normalize(cell.text)
                    values.append(text)
                    if text:
                        cells.append(
                            self._store_text(
                                artifact,
                                text,
                                SourceLocator(
                                    kind="docx_table_cell",
                                    value=(
                                        f"table:{table_index}:row:{row_index}:cell:{cell_index}"
                                    ),
                                    artifact_id=artifact.id,
                                    table=str(table_index),
                                    cell=f"R{row_index}C{cell_index}",
                                ),
                            )
                        )
                rows.append(values)
            text = "\n".join("\t".join(row) for row in rows)
            stored = self._artifact_store.put_bytes(
                text.encode("utf-8"),
                media_type="text/tab-separated-values; charset=utf-8",
                artifact_id=artifact.id,
                source_snapshot_hash=artifact.content_hash,
                sensitivity=artifact.sensitivity,
            )
            parsed.append(
                ParsedTable(
                    locator=SourceLocator(
                        kind="docx_table",
                        value=f"table:{table_index}",
                        artifact_id=artifact.id,
                        table=str(table_index),
                    ),
                    text_ref=stored.content_hash,
                    content_hash=stored.content_hash,
                    char_count=len(text),
                    row_count=len(rows),
                    column_count=max((len(row) for row in rows), default=0),
                    text_blocks=cells,
                    confidence=Decimal("1"),
                )
            )
        return parsed

    def _store_text(self, artifact: Artifact, text: str, locator: SourceLocator) -> TextBlock:
        stored = self._artifact_store.put_bytes(
            text.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            artifact_id=artifact.id,
            source_snapshot_hash=artifact.content_hash,
            sensitivity=artifact.sensitivity,
        )
        return TextBlock(
            text_ref=stored.content_hash,
            content_hash=stored.content_hash,
            char_count=len(text),
            locator=locator,
            confidence=Decimal("1"),
            verification_status="verified",
        )


def _normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()
