from __future__ import annotations

from datetime import UTC, datetime
from hashlib import sha256
from io import BytesIO
from pathlib import Path
import traceback
from typing import Callable
from uuid import NAMESPACE_URL, UUID, uuid4, uuid5
import duckdb

from docx import Document
from openpyxl import Workbook
import pytest

from due_diligence_agent.adapters.startup.profile_fragment_inventory import (
    StartupProfileFragmentInventoryIntegrityError,
    _bounded_fragment_text,
)
from due_diligence_agent.bootstrap.container import (
    build_deterministic_startup_analysis_composer,
    build_local_repositories,
    build_startup_analysis_composer,
)
from due_diligence_agent.domain.approvals.startup_disclosure import ClassifiedDisclosureSnapshot
from due_diligence_agent.domain.artifacts.models import Artifact
from due_diligence_agent.domain.cases.models import DueDiligenceCase
from due_diligence_agent.domain.common import AnalysisMode, CaseStatus, ContradictionStatus, SensitivityClass
from due_diligence_agent.domain.startup.profile import StartupProfileFieldName, StartupProfileFieldStatus
from due_diligence_agent.ports.startup_profile_extraction import (
    MAX_FRAGMENTS,
    MAX_TOTAL_FRAGMENT_CHARS,
)


CASE_ID = UUID("aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa")


@pytest.mark.parametrize(
    ("source", "payload_factory", "mime_type"),
    [
        (
            "metrics.csv",
            lambda: b"Metric,2026\r\nARR,1200\r\n",
            "text/csv",
        ),
        (
            "metrics.xlsx",
            lambda: _xlsx_bytes(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    ],
)
def test_default_startup_parsing_composer_parses_spreadsheets_without_flattening(
    source: str,
    payload_factory: Callable[[], bytes],
    mime_type: str,
) -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = payload_factory()
    artifact = _artifact(payload, source, mime_type=mime_type)
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)

    parsed = parser.parse(artifact)

    assert parsed.kind == "spreadsheet"
    assert parsed.status == "parsed"
    assert parsed.spreadsheet is not None
    assert parser.documents([str(artifact.id)]) == []

    if source.endswith(".xlsx"):
        database_path = data_dir / "startup-normalized.duckdb"
        assert database_path.exists()
        with duckdb.connect(str(database_path), read_only=True) as connection:
            rows = connection.execute(
                "SELECT table_name FROM normalized_table_snapshots ORDER BY table_name"
            ).fetchall()
        assert rows == [("Financials",)]


def test_default_startup_parsing_composer_keeps_document_text_blocks(
) -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _docx_bytes()
    artifact = _artifact(
        payload,
        "governance.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)

    parsed = parser.parse(artifact)

    assert parsed.kind == "document"
    assert parsed.document is not None
    assert parser.documents([str(artifact.id)]) == [parsed.document]
    assert {block.locator.value for block in parser.text_blocks([str(artifact.id)])} >= {
        "paragraph:1"
    }


def test_startup_text_block_and_disclosure_order_ignore_case_bound_artifact_ids() -> None:
    case_a = UUID("aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaa1")
    case_b = UUID("bbbbbbbb-bbbb-bbbb-bbbb-bbbbbbbbbbb2")
    low_id = UUID("11111111-1111-1111-1111-111111111111")
    high_id = UUID("eeeeeeee-eeee-eeee-eeee-eeeeeeeeeeee")
    first_payload = _single_paragraph_docx_bytes("Problem: stable ordering across case identity")
    second_payload = _single_paragraph_docx_bytes("Solution: deterministic disclosure snapshot")

    first = _text_order_projection(
        case_id=case_a,
        artifacts=((low_id, first_payload), (high_id, second_payload)),
    )
    second = _text_order_projection(
        case_id=case_b,
        artifacts=((high_id, first_payload), (low_id, second_payload)),
    )

    assert first == second


def test_startup_parsing_composer_reloads_persisted_document_after_restart() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _docx_bytes()
    artifact = _artifact(
        payload,
        "restart-governance.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)

    parsed = parser.parse(artifact)

    restarted = build_startup_analysis_composer(data_dir)
    restarted_parser = restarted._dependencies.parser

    assert restarted_parser.documents([str(parsed.artifact_id)]) == [parsed.document]
    assert {
        block.locator.value for block in restarted_parser.text_blocks([str(parsed.artifact_id)])
    } >= {"paragraph:1"}
    evidence = restarted._dependencies.evidence.extract(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
    )
    assert len(evidence["evidence_fact_ids"]) >= 1


def test_startup_data_room_resolves_safe_source_refs_and_reuses_existing_artifacts() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    case_id = str(CASE_ID)
    inbox_case_dir = data_dir / "inbox" / case_id
    inbox_case_dir.mkdir(parents=True)
    payload = _profile_docx_bytes()
    content_hash = sha256(payload).hexdigest()
    (inbox_case_dir / "doc-0001.docx").write_bytes(payload)
    source_refs = [
        {
            "document_id": "doc-0001",
            "private_name": "doc-0001.docx",
            "content_sha256": content_hash,
        }
    ]

    first = service._dependencies.data_room.ingest(
        case_id=case_id,
        source_refs=source_refs,
        data_revision=1,
    )
    second = service._dependencies.data_room.ingest(
        case_id=case_id,
        source_refs=source_refs,
        data_revision=1,
    )

    assert first["inventory_id"] == second["inventory_id"]
    assert first["artifact_ids"] == second["artifact_ids"]
    assert first["quarantine"] == []
    assert all(str(inbox_case_dir) not in value for value in first["artifact_ids"])


def test_deterministic_startup_composer_ingests_shared_inbox_refs_with_isolated_storage() -> None:
    live_data_dir = _workspace_tmp_dir() / "startup-api"
    deterministic_data_dir = live_data_dir / "deterministic"
    shared_inbox_root = live_data_dir / "inbox"
    service = build_deterministic_startup_analysis_composer(
        deterministic_data_dir,
        inbox_root=shared_inbox_root,
    )
    case_id = str(CASE_ID)
    inbox_case_dir = shared_inbox_root / case_id
    inbox_case_dir.mkdir(parents=True)
    payload = _profile_docx_bytes()
    content_hash = sha256(payload).hexdigest()
    (inbox_case_dir / "doc-0001.docx").write_bytes(payload)

    inventory = service._dependencies.data_room.ingest(
        case_id=case_id,
        source_refs=[
            {
                "document_id": "doc-0001",
                "private_name": "doc-0001.docx",
                "content_sha256": content_hash,
            }
        ],
        data_revision=1,
    )

    repositories = build_local_repositories(deterministic_data_dir / "startup-metadata.sqlite3")
    artifacts = repositories.artifact_repository.list_for_case(CASE_ID)
    assert inventory["quarantine"] == []
    assert inventory["artifact_ids"] == [str(artifacts[0].id)]
    assert artifacts[0].source_snapshot_hash == content_hash
    assert (deterministic_data_dir / "startup-artifacts").exists()
    assert not (live_data_dir / "startup-artifacts").exists()
    assert not (live_data_dir / "startup-metadata.sqlite3").exists()


def test_startup_composer_reloads_redacted_profile_fragments_and_spreadsheet_facts() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case(data_revision=2))
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    document_payload = _profile_docx_bytes()
    spreadsheet_payload = _evidence_xlsx_bytes()
    document_artifact = _artifact(
        document_payload,
        "startup-profile.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    spreadsheet_artifact = _artifact(
        spreadsheet_payload,
        "startup-metrics.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    for artifact, payload in (
        (document_artifact, document_payload),
        (spreadsheet_artifact, spreadsheet_payload),
    ):
        repositories.artifact_repository.add(artifact)
        _persist(artifact_store, artifact, payload)

    parsed_document = parser.parse(document_artifact)
    parsed_spreadsheet = parser.parse(spreadsheet_artifact)
    service._dependencies.evidence.extract(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed_document.artifact_id), str(parsed_spreadsheet.artifact_id)],
    )
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed_document.artifact_id), str(parsed_spreadsheet.artifact_id)],
        data_revision=2,
    )
    snapshot = privacy["snapshot"]
    assert snapshot.data_revision == 2
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": snapshot})

    restarted = build_startup_analysis_composer(data_dir)
    fragments = restarted._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 2)
    evidence_facts = build_local_repositories(
        data_dir / "startup-metadata.sqlite3"
    ).evidence_repository.list_for_case(CASE_ID)
    profile = restarted._dependencies.profile.build_primary(CASE_ID)

    assert [fragment.artifact_id for fragment in fragments] == [document_artifact.id] * 3
    assert [fragment.text for fragment in fragments] == [
        "Startup Name: Atlas AI",
        "Problem: Founders do not know what go-to-market risks they missed",
        "Solution: Automated startup readiness analysis",
    ]
    assert [(fact.name, str(fact.value), fact.unit, fact.period) for fact in evidence_facts if fact.name == "revenue"] == [
        ("revenue", "1250.5", "USD", "2025")
    ]
    assert profile.data_revision == 2
    assert profile.fields["startup_name"].values == ("Atlas AI",)
    assert profile.fields["traction"].values == ("revenue: 1250.5 USD 2025",)


def test_startup_evidence_port_materializes_contradictions_from_parsed_csv_and_xlsx_scalars() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    csv_payload = b"Metric,Value,Unit,Period\r\nOrders,720,count,2026-07\r\n"
    xlsx_payload = _scalar_evidence_xlsx_bytes(
        metric="Orders",
        value=680,
        unit="count",
        period="2026-07",
    )
    csv_artifact = _artifact(csv_payload, "orders-source-a.csv", mime_type="text/csv")
    xlsx_artifact = _artifact(
        xlsx_payload,
        "orders-source-b.xlsx",
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    for artifact, payload in ((csv_artifact, csv_payload), (xlsx_artifact, xlsx_payload)):
        repositories.artifact_repository.add(artifact)
        _persist(artifact_store, artifact, payload)

    parsed_csv = parser.parse(csv_artifact)
    parsed_xlsx = parser.parse(xlsx_artifact)
    result = service._dependencies.evidence.extract(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed_csv.artifact_id), str(parsed_xlsx.artifact_id)],
    )
    persisted_facts = repositories.evidence_repository.list_for_case(CASE_ID)
    persisted_contradictions = repositories.contradiction_repository.list_for_case(CASE_ID)

    assert len(persisted_contradictions) == 1
    assert sorted((fact.name, str(fact.value), fact.unit, fact.period) for fact in persisted_facts) == [
        ("orders", "680", "count", "2026-07"),
        ("orders", "720", "count", "2026-07"),
    ]
    assert set(result["evidence_fact_ids"]) == {str(fact.id) for fact in persisted_facts}
    assert result["contradiction_ids"] == [str(persisted_contradictions[0].id)]
    assert set(persisted_contradictions[0].fact_ids) == {fact.id for fact in persisted_facts}


def test_startup_profile_fragment_inventory_bounds_long_redacted_blocks() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _single_paragraph_docx_bytes("Problem: " + ("market risk " * 120))
    artifact = _artifact(
        payload,
        "long-profile.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": privacy["snapshot"]})

    fragments = service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)

    assert fragments
    assert all(len(fragment.text) <= 800 for fragment in fragments)


def test_startup_profile_fragment_inventory_samples_aligned_snapshot_over_request_limit() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _multi_paragraph_docx_bytes(MAX_FRAGMENTS + 7)
    artifact = _artifact(
        payload,
        "large-profile.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    snapshot = privacy["snapshot"]
    service._dependencies.workflow_store.save(
        str(CASE_ID),
        {"disclosure_snapshot": snapshot},
    )

    fragments = service._dependencies.profile._fragment_inventory.list_for_case_revision(
        CASE_ID,
        1,
    )

    assert len(snapshot.redacted_fragment_ids) == MAX_FRAGMENTS + 7
    assert len(fragments) == MAX_FRAGMENTS
    assert sum(len(fragment.text) for fragment in fragments) <= MAX_TOTAL_FRAGMENT_CHARS
    selected_fragment_ids = {fragment.fragment_id for fragment in fragments}
    assert snapshot.redacted_fragment_ids[0] in selected_fragment_ids
    assert snapshot.redacted_fragment_ids[-1] in selected_fragment_ids
    restarted = build_startup_analysis_composer(data_dir)
    restarted_fragments = (
        restarted._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)
    )
    assert tuple(fragment.fragment_id for fragment in restarted_fragments) == tuple(
        fragment.fragment_id for fragment in fragments
    )


def test_startup_profile_fragment_inventory_prioritizes_profile_and_conflict_blocks_over_filler() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    high_value = [
        "NomadFlow AI",
        "NomadFlow AI разрабатывает облачную платформу управления запасами, закупками и маршрутизацией.",
        "Приоритетные сегменты: FMCG-дистрибьюторы, региональный retail и фармацевтическая дистрибуция.",
        "Тарифы: Starter 240 000 ₸/мес, Growth 690 000 ₸/мес, Enterprise 1 900 000 ₸/мес.",
        "MRR CONTRADICTION CRM 28,6 млн ₸; invoices 27,9 млн ₸",
        "Customers CONTRADICTION CRM 31; invoices 29",
        "Gross margin CONTRADICTION operational 74%; fully-loaded 70%",
        "CAC payback CONTRADICTION reported 4.3 months; recalculated 5.5 months",
    ]
    filler = [
        f"Appendix {index:02d}: process note without startup profile facts"
        for index in range(MAX_FRAGMENTS + 12)
    ]
    payload = _paragraphs_docx_bytes([*filler[:12], *high_value, *filler[12:]])
    artifact = _artifact(
        payload,
        "nomadflow-long-profile.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": privacy["snapshot"]})

    fragments = service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)
    selected_text = "\n".join(fragment.text for fragment in fragments)

    assert len(fragments) == MAX_FRAGMENTS
    for expected in (
        "NomadFlow AI",
        "облачную платформу управления запасами",
        "FMCG-дистрибьюторы",
        "Starter 240 000 ₸/мес",
        "MRR CONTRADICTION CRM 28,6 млн ₸",
        "Customers CONTRADICTION CRM 31",
        "Gross margin CONTRADICTION operational 74%",
        "CAC payback CONTRADICTION reported 4.3 months",
    ):
        assert expected in selected_text


def test_startup_profile_fragment_inventory_selects_canonical_nomadflow_pdf_profile_signals() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = (Path("output") / "pdf" / "nomadflow_ai_startup_test_business_plan_ru.pdf").read_bytes()
    artifact = _artifact(payload, "nomadflow-business-plan.pdf", mime_type="application/pdf")
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": privacy["snapshot"]})

    fragments = service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)
    restarted_fragments = build_startup_analysis_composer(
        data_dir
    )._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)
    selected_text = "\n".join(fragment.text for fragment in fragments)

    assert len(fragments) <= MAX_FRAGMENTS
    assert sum(len(fragment.text) for fragment in fragments) <= MAX_TOTAL_FRAGMENT_CHARS
    assert tuple(fragment.fragment_id for fragment in restarted_fragments) == tuple(
        fragment.fragment_id for fragment in fragments
    )
    assert any(fragment.page == 1 for fragment in fragments)
    for expected in (
        "NomadFlow AI",
        "облачную платформу управления запасами, планирования закупок и маршрутизации доставок",
        "Стадия Seed",
        "География Казахстан; далее Узбекистан и Кыргызстан",
        "Starter 240 тыс. ₸",
        "Growth 690 тыс. ₸",
        "Enterprise от 1,9 млн ₸",
        "MRR CONTRADICTION CRM 28,6 млн ₸; invoices 27,9 млн ₸",
        "Клиенты CONTRADICTION 31 в CRM; 29 invoiced",
        "Gross margin CONTRADICTION 74% operational; 70% fully loaded",
        "CAC payback CONTRADICTION 4,3 мес. заявлено; 5,5 мес. пересчет",
        "net burn за последние три месяца: 22,4 млн ₸",
        "runway: 7,8 месяца",
    ):
        assert expected in selected_text


def test_startup_primary_profile_extracts_core_fields_from_canonical_nomadflow_pdf() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    payload = (Path("output") / "pdf" / "nomadflow_ai_startup_test_business_plan_ru.pdf").read_bytes()
    artifact = _artifact(payload, "nomadflow-business-plan.pdf", mime_type="application/pdf")
    repositories.artifact_repository.add(artifact)
    _persist(parser._artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": privacy["snapshot"]})

    profile = service._dependencies.profile.build_primary(CASE_ID)

    core_fields = (
        StartupProfileFieldName.STARTUP_NAME,
        StartupProfileFieldName.ONE_LINE_DESCRIPTION,
        StartupProfileFieldName.PROBLEM,
        StartupProfileFieldName.ICP,
        StartupProfileFieldName.PRICING_REVENUE_MODEL,
        StartupProfileFieldName.STAGE,
    )
    for field_name in core_fields:
        field = profile.fields[field_name.value]
        assert field.status is StartupProfileFieldStatus.SOURCE_FACT
        assert field.values
        assert field.evidence_refs
        assert all(ref.artifact_id == artifact.id for ref in field.evidence_refs)
        assert all(ref.artifact_hash == f"sha256:{artifact.content_hash}" for ref in field.evidence_refs)
        assert all(ref.locator_hash for ref in field.evidence_refs)

    assert "NomadFlow AI" in profile.fields["startup_name"].values
    assert any("запасами" in value and "маршрутизации" in value for value in profile.fields["one_line_description"].values)
    assert any("ручн" in value.casefold() or "потер" in value.casefold() for value in profile.fields["problem"].values)
    assert any(
        "дистриб" in value.casefold() or "retail" in value.casefold() or "фарма" in value.casefold()
        for value in profile.fields["icp"].values
    )
    assert any("Seed" in value for value in profile.fields["stage"].values)
    assert any("Starter" in value for value in profile.fields["pricing_revenue_model"].values)
    assert any("Growth" in value for value in profile.fields["pricing_revenue_model"].values)
    dumped_profile = profile.model_dump_json()
    assert "C:\\Users" not in dumped_profile
    assert "D:\\Agents" not in dumped_profile
    assert "Describe the product" not in dumped_profile
    assert "Опишите продукт в одну строку" not in dumped_profile


def test_startup_evidence_materializes_four_canonical_nomadflow_pdf_table_contradictions() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    payload = (Path("output") / "pdf" / "nomadflow_ai_startup_test_business_plan_ru.pdf").read_bytes()
    artifact = _artifact(payload, "nomadflow-business-plan.pdf", mime_type="application/pdf")
    repositories.artifact_repository.add(artifact)
    _persist(parser._artifact_store, artifact, payload)
    parsed = parser.parse(artifact)

    result = service._dependencies.evidence.extract(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
    )
    contradictions = repositories.contradiction_repository.list_for_case(CASE_ID)
    table_blocks = [
        block
        for block in parser.text_blocks([str(parsed.artifact_id)])
        if block.locator.kind == "pdf_table"
    ]
    dumped = "\n".join(item.model_dump_json() for item in contradictions)

    assert len(result["contradiction_ids"]) == 4
    assert len(contradictions) == 4
    assert all(item.status is ContradictionStatus.OPEN for item in contradictions)
    assert all(item.conflict_type == "explicit_source_conflict_signal" for item in contradictions)
    assert any(block.locator.page == 14 and block.locator.table == "1" for block in table_blocks)
    facts_by_id = {
        fact.id: fact for fact in repositories.evidence_repository.list_for_case(CASE_ID)
    }
    startup_metric_values = {
        (fact.name, str(fact.value), fact.unit)
        for fact in facts_by_id.values()
        if fact.metadata.get("startup_claim_id") is not None
    }
    text_by_hash = {
        block.content_hash: parser.text_for_block(block)
        for block in parser.text_blocks([str(parsed.artifact_id)])
    }
    contradiction_source_texts = [
        text_by_hash[facts_by_id[fact_id].supporting_text_hash]
        for contradiction in contradictions
        for fact_id in contradiction.fact_ids
    ]
    assert _canonical_pdf_contradiction_topics(contradiction_source_texts) == {
        "mrr",
        "customer_count",
        "gross_margin",
        "cac_payback",
    }
    assert ("monthly_recurring_revenue", "28600000.0", "KZT") in startup_metric_values
    assert ("monthly_recurring_revenue", "27900000.0", "KZT") in startup_metric_values
    assert ("gross_margin", "74", "percent") in startup_metric_values
    assert ("gross_margin", "70", "percent") in startup_metric_values
    assert ("monthly_net_burn", "22400000.0", "KZT/month") in startup_metric_values
    assert ("runway", "7.8", "months") in startup_metric_values
    assert ("gross_margin", "5", "percent") not in startup_metric_values
    assert ("gross_margin", "28", "percent") not in startup_metric_values
    assert ("runway", "18", "months") not in startup_metric_values
    for private_value in ("28,6 млн", "27,9 млн", "31 в CRM", "29 invoiced", "74%", "70%", "4,3", "5,5"):
        assert private_value not in dumped


def _canonical_pdf_contradiction_topics(texts: list[str]) -> set[str]:
    topics: set[str] = set()
    for text in texts:
        for line in text.splitlines():
            if "CONTRADICTION" not in line:
                continue
            if "MRR" in line and "28,6" in line and "27,9" in line:
                topics.add("mrr")
            if "Клиент" in line and "31 в CRM" in line and "29 invoiced" in line:
                topics.add("customer_count")
            if "Gross margin" in line and "74%" in line and "70%" in line:
                topics.add("gross_margin")
            if "CAC payback" in line and "4,3" in line and "5,5" in line:
                topics.add("cac_payback")
    return topics


def test_startup_profile_fragment_inventory_keeps_relevant_tail_lines_when_table_fragment_is_long() -> None:
    long_table = "\n".join(
        [
            "Альтернатива Сильная сторона Ограничение для ICP Ответ NomadFlow",
            *[
                f"Competitor row {index:02d} локализация внедрение ограничение"
                for index in range(30)
            ],
            "Тариф Ежемесячно Разовый запуск Включено",
            "Starter 240 тыс. ₸ 900 тыс. ₸ 1 склад; запасы; базовые алерты",
            "Growth 690 тыс. ₸ 2,4 млн ₸ до 5 складов; спрос; маршруты",
            "Enterprise от 1,9 млн ₸ от 6 млн ₸ multi-entity; SLA; private connectors",
        ]
    )

    excerpt = _bounded_fragment_text(long_table, max_chars=500)

    assert "Competitor row 00" not in excerpt
    assert "Тариф Ежемесячно" in excerpt
    assert "Starter 240 тыс. ₸" in excerpt
    assert "Growth 690 тыс. ₸" in excerpt
    assert "Enterprise от 1,9 млн ₸" in excerpt


def test_startup_profile_fragment_inventory_keeps_idea_brief_name_when_fragment_is_long() -> None:
    long_idea_brief = "\n".join(
        [
            "Founder idea brief: CareLoop Recall",
            "",
            "Concept:",
            (
                "CareLoop Recall is a follow-up coordination service for outpatient clinics "
                "that lose patient continuity after visits because reminders, lab callbacks, "
                "and chronic-care check-ins are scattered across phone calls and manual notes."
            ),
            "",
            "Buyer:",
            (
                "Clinic administrators and medical directors who need a reliable follow-up "
                "process without replacing the existing appointment system."
            ),
            "",
            "Known gaps:",
            (
                "No uploaded document provides revenue, expense, burn, cash balance, customer "
                "count, ARR, MRR, or cohort metrics. The case is idea-only and should keep "
                "planning scenarios separate from evidence-backed facts."
            ),
        ]
    )

    excerpt = _bounded_fragment_text(long_idea_brief, max_chars=500)

    assert "Founder idea brief: CareLoop Recall" in excerpt
    assert "CareLoop Recall is a follow-up coordination service" in excerpt
    assert "No uploaded document provides revenue" in excerpt


def test_startup_profile_fragment_inventory_preserves_empty_aligned_snapshot() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    service._dependencies.workflow_store.save(
        str(CASE_ID),
        {"disclosure_snapshot": _snapshot(fragment_ids=(), text_refs=())},
    )

    fragments = service._dependencies.profile._fragment_inventory.list_for_case_revision(
        CASE_ID,
        1,
    )

    assert fragments == ()


def test_startup_profile_fragment_inventory_fails_closed_on_revision_mismatch() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case(data_revision=2))
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _profile_docx_bytes()
    artifact = _artifact(
        payload,
        "startup-profile.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parsed = parser.parse(artifact)
    privacy = service._dependencies.privacy.classify_redact(
        case_id=str(CASE_ID),
        parsed_artifact_ids=[str(parsed.artifact_id)],
        data_revision=1,
    )
    service._dependencies.workflow_store.save(str(CASE_ID), {"disclosure_snapshot": privacy["snapshot"]})

    with pytest.raises(StartupProfileFragmentInventoryIntegrityError) as exc_info:
        service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 2)

    assert "startup_disclosure_snapshot_revision_mismatch" in str(exc_info.value)


def test_startup_profile_fragment_inventory_normalizes_missing_minimized_fragment_file() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case())
    parser = service._dependencies.parser
    artifact_store = parser._artifact_store
    payload = _single_paragraph_docx_bytes("Problem: founders need safer diligence")
    artifact = _artifact(
        payload,
        "missing-minimized-fragment.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    repositories.artifact_repository.add(artifact)
    _persist(artifact_store, artifact, payload)
    parser.parse(artifact)
    block = parser.text_blocks_for_case(CASE_ID)[0]
    missing_text_ref = "0" * 64
    fragment_id = uuid5(NAMESPACE_URL, f"{block.content_hash}:{missing_text_ref}")
    service._dependencies.workflow_store.save(
        str(CASE_ID),
        {"disclosure_snapshot": _snapshot(fragment_ids=(fragment_id,), text_refs=(missing_text_ref,))},
    )

    with pytest.raises(StartupProfileFragmentInventoryIntegrityError) as exc_info:
        service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)

    assert exc_info.value.stable_error_code == "STARTUP_PROFILE_FRAGMENT_INVENTORY_INTEGRITY_ERROR"
    assert str(exc_info.value) == "startup_profile_fragment_inventory_integrity_error"
    assert exc_info.value.__cause__ is None
    public_traceback = "".join(
        traceback.format_exception(exc_info.type, exc_info.value, exc_info.tb)
    )
    assert str(_workspace_tmp_root()) not in public_traceback
    assert "startup-artifacts" not in public_traceback
    assert "objects" not in public_traceback
    assert missing_text_ref not in public_traceback


def test_startup_profile_fragment_inventory_rejects_empty_fragment_ids_with_minimized_refs() -> None:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    service._dependencies.workflow_store.save(
        str(CASE_ID),
        {"disclosure_snapshot": _snapshot(fragment_ids=(), text_refs=("1" * 64,))},
    )

    with pytest.raises(StartupProfileFragmentInventoryIntegrityError) as exc_info:
        service._dependencies.profile._fragment_inventory.list_for_case_revision(CASE_ID, 1)

    assert str(exc_info.value) == "startup_profile_fragment_inventory_snapshot_mismatch"


def _case(*, data_revision: int = 1) -> DueDiligenceCase:
    now = datetime(2026, 8, 13, tzinfo=UTC)
    return DueDiligenceCase(
        case_id=CASE_ID,
        mode=AnalysisMode.STARTUP,
        entity_name="Startup",
        entity_identifier="startup",
        jurisdiction="KZ",
        scope=("startup_due_diligence",),
        as_of=now,
        base_currency="USD",
        privacy_policy="startup-local@1",
        budget_policy="startup-local@1",
        status=CaseStatus.CREATED,
        sensitivity=SensitivityClass.RESTRICTED,
        created_at=now,
        updated_at=now,
        workflow_version="startup-local@1",
        data_revision=data_revision,
    )


def _snapshot(
    *,
    fragment_ids: tuple[UUID, ...],
    text_refs: tuple[str, ...],
    data_revision: int = 1,
) -> ClassifiedDisclosureSnapshot:
    return ClassifiedDisclosureSnapshot(
        case_id=CASE_ID,
        detected_classes=frozenset({SensitivityClass.RESTRICTED}),
        overall_class=SensitivityClass.RESTRICTED,
        redaction_policy_version="startup-local@1",
        egress_policy_version="startup-local@1",
        data_revision=data_revision,
        content_hash="f" * 64,
        artifact_counts={"document": 1},
        mime_counts={"docx": 1},
        category_counts={"profile": 1},
        redacted_fragment_ids=fragment_ids,
        minimized_fragment_refs=text_refs,
        destination="startup_profile_extraction",
    )


def _artifact(payload: bytes, source: str, *, mime_type: str) -> Artifact:
    digest = sha256(payload).hexdigest()
    return Artifact(
        id=uuid4(),
        case_id=CASE_ID,
        content_hash=digest,
        mime_type=mime_type,
        source=source,
        retrieved_at=datetime.now(UTC),
        source_snapshot_hash=digest,
        sensitivity=SensitivityClass.RESTRICTED,
    )


def _persist(artifact_store: object, artifact: Artifact, payload: bytes) -> None:
    artifact_store.put_bytes(
        payload,
        media_type=artifact.mime_type,
        artifact_id=artifact.id,
        source_snapshot_hash=artifact.source_snapshot_hash,
        sensitivity=artifact.sensitivity,
    )


def _text_order_projection(
    *,
    case_id: UUID,
    artifacts: tuple[tuple[UUID, bytes], tuple[UUID, bytes]],
) -> tuple[list[str], tuple[str, ...], tuple[UUID, ...], str]:
    data_dir = _workspace_tmp_dir()
    service = build_startup_analysis_composer(data_dir)
    repositories = build_local_repositories(data_dir / "startup-metadata.sqlite3")
    repositories.case_repository.add(_case().model_copy(update={"case_id": case_id}))
    parser = service._dependencies.parser
    parsed_ids: list[str] = []
    for index, (artifact_id, payload) in enumerate(artifacts, start=1):
        digest = sha256(payload).hexdigest()
        artifact = Artifact(
            id=artifact_id,
            case_id=case_id,
            content_hash=digest,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source=f"document-{index}.docx",
            retrieved_at=datetime.now(UTC),
            source_snapshot_hash=digest,
            sensitivity=SensitivityClass.RESTRICTED,
        )
        repositories.artifact_repository.add(artifact)
        _persist(parser._artifact_store, artifact, payload)
        parsed_ids.append(str(parser.parse(artifact).artifact_id))

    blocks = parser.text_blocks(parsed_ids)
    snapshot = service._dependencies.privacy.classify_redact(
        case_id=str(case_id),
        parsed_artifact_ids=parsed_ids,
        data_revision=1,
    )["snapshot"]
    return (
        [block.content_hash for block in blocks],
        snapshot.minimized_fragment_refs,
        snapshot.redacted_fragment_ids,
        snapshot.content_hash,
    )


def _xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Financials"
    sheet["A1"] = "Metric"
    sheet["B1"] = "2026"
    sheet["A2"] = "ARR"
    sheet["B2"] = 1200
    sheet["B2"].number_format = "$#,##0"
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _evidence_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Evidence"
    sheet["A1"] = "Metric"
    sheet["B1"] = "2025"
    sheet["A2"] = "Revenue"
    sheet["B2"] = 1250.50
    sheet["B2"].number_format = "$#,##0.00"
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _scalar_evidence_xlsx_bytes(
    *,
    metric: str,
    value: int | float,
    unit: str,
    period: str,
) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Evidence"
    sheet["A1"] = "Metric"
    sheet["B1"] = "Value"
    sheet["C1"] = "Unit"
    sheet["D1"] = "Period"
    sheet["A2"] = metric
    sheet["B2"] = value
    sheet["C2"] = unit
    sheet["D2"] = period
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Board composition")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Director"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Ada"
    table.cell(1, 1).text = "Chair"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _profile_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Startup Name: Atlas AI")
    document.add_paragraph("Problem: Founders do not know what go-to-market risks they missed")
    document.add_paragraph("Solution: Automated startup readiness analysis")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _single_paragraph_docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _multi_paragraph_docx_bytes(count: int) -> bytes:
    document = Document()
    for index in range(count):
        document.add_paragraph(
            f"Section {index:03d}: problem solution market competition pricing GTM risk"
        )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _paragraphs_docx_bytes(paragraphs: list[str]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _workspace_tmp_root() -> Path:
    root = Path("artifacts") / "test-runs" / "startup-parsing-composition"
    root.mkdir(parents=True, exist_ok=True)
    return root


def _workspace_tmp_dir() -> Path:
    root = _workspace_tmp_root() / f"run-{uuid4().hex}"
    root.mkdir(parents=True, exist_ok=False)
    return root
