from __future__ import annotations

from datetime import UTC, datetime
from decimal import Decimal
import importlib
from io import BytesIO
from hashlib import sha256
import json
import os
from pathlib import Path
import socket
import subprocess
import sys
from types import SimpleNamespace
from uuid import UUID, uuid4

import fitz
from docx import Document
from PIL import Image, PngImagePlugin
import pytest

from due_diligence_agent.adapters.documents.docling_parser import DoclingDocumentParser
from due_diligence_agent.adapters.documents.image_ocr import (
    ImageSafetyLimits,
    OcrRegion,
    TesseractOcrAdapter,
)
from due_diligence_agent.adapters.documents.no_network_guard import (
    NoNetworkGuard,
    NoNetworkViolation,
)
from due_diligence_agent.adapters.local_storage.artifact_store import LocalArtifactStore
from due_diligence_agent.application.services.startup_parsing_service import (
    StartupParsingService,
)
from due_diligence_agent.domain.artifacts.models import Artifact
from due_diligence_agent.domain.common import SensitivityClass


def test_pdf_text_has_stable_page_and_coordinate_locator(tmp_path: Path) -> None:
    payload = _pdf_with_text("Annual recurring revenue")
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "renamed.bin")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    block = parsed.pages[0].text_blocks[0]
    assert parsed.status == "parsed"
    assert block.locator.kind == "pdf_text_block"
    assert block.locator.page == 1
    assert block.locator.value.startswith("page:1:bbox:")
    assert block.char_count == 24
    assert store.read_bytes(block.text_ref).decode() == "Annual recurring revenue"


def test_pdf_table_uses_pdfplumber_fallback_with_stable_locator(tmp_path: Path) -> None:
    payload = _pdf_with_table()
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "table.pdf")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert len(parsed.tables) == 1
    assert parsed.tables[0].locator.kind == "pdf_table"
    assert parsed.tables[0].locator.value == "page:1:table:1"
    assert parsed.tables[0].row_count == 2
    assert parsed.tables[0].column_count == 2
    assert [block.locator.value for block in parsed.tables[0].text_blocks] == [
        "page:1:table:1",
    ]
    assert parsed.tables[0].text_blocks[0].locator.page == 1
    assert parsed.tables[0].text_blocks[0].locator.table == "1"
    assert store.read_bytes(parsed.tables[0].text_ref) == b"Metric\tValue\nARR\t1200"


def test_docx_paragraph_and_table_have_stable_locators(tmp_path: Path) -> None:
    payload = _docx_bytes()
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "governance.docx")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.text_blocks[0].locator.value == "paragraph:1"
    assert parsed.tables[0].locator.value == "table:1"
    assert parsed.tables[0].text_blocks[0].locator.value == "table:1:row:1:cell:1"
    assert store.read_bytes(parsed.text_blocks[0].text_ref) == b"Board composition"


def test_image_is_validated_and_normalized_before_ocr(tmp_path: Path) -> None:
    seen: list[tuple[str, tuple[int, int]]] = []

    def recognize(image: Image.Image) -> list[OcrRegion]:
        seen.append((image.mode, image.size))
        return [OcrRegion(text="Runway 18 months", confidence=Decimal("0.95"), bbox=(1, 2, 8, 9))]

    service, store = _service(tmp_path, ocr=TesseractOcrAdapter(recognizer=recognize))
    payload = _image_bytes(mode="P")
    artifact = _artifact(payload, "scan.dat")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parsed"
    assert seen == [("RGB", (12, 10))]
    assert parsed.pages[0].text_blocks[0].locator.value == "page:1:bbox:1,2,8,9"


def test_scanned_value_is_not_verified_when_ocr_confidence_is_low(tmp_path: Path) -> None:
    def recognize(_image: Image.Image) -> list[OcrRegion]:
        return [OcrRegion(text="MRR maybe 42", confidence=Decimal("0.61"), bbox=(0, 0, 9, 9))]

    service, store = _service(tmp_path, ocr=TesseractOcrAdapter(recognizer=recognize))
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    value = parsed.pages[0].text_blocks[0]
    assert value.locator.kind == "image_region"
    assert value.confidence < Decimal("0.80")
    assert value.verification_status == "needs_review"


def test_missing_tesseract_binary_returns_typed_parser_unavailable(tmp_path: Path) -> None:
    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(binary_probe=lambda: None),
    )
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parser_unavailable"
    assert parsed.error_code == "ocr_binary_missing"
    assert parsed.text_blocks == []


def test_binary_present_with_missing_python_adapter_is_typed_unavailable(tmp_path: Path) -> None:
    def missing_adapter(_name: str) -> object:
        raise ModuleNotFoundError("SECRET local module search path")

    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(
            binary_probe=lambda: "C:/local/tesseract.exe",
            version_runner=lambda _command: subprocess.CompletedProcess(
                args=[], returncode=0, stdout="tesseract 5.5.0", stderr=""
            ),
            module_loader=missing_adapter,
        ),
    )
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parser_unavailable"
    assert parsed.error_code == "ocr_python_adapter_missing"
    assert "SECRET" not in repr(parsed)


def test_broken_python_adapter_import_is_typed_unusable_without_raw_leak(
    tmp_path: Path,
) -> None:
    def broken_adapter(_name: str) -> object:
        raise RuntimeError("SECRET broken optional install")

    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(
            binary_probe=lambda: "C:/local/tesseract.exe",
            version_runner=lambda _command: subprocess.CompletedProcess(
                args=[], returncode=0, stdout="tesseract 5.5.0", stderr=""
            ),
            module_loader=broken_adapter,
        ),
    )
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parser_unavailable"
    assert parsed.error_code == "ocr_python_adapter_unusable"
    assert "SECRET" not in repr(parsed)
    assert "SECRET" not in parsed.model_dump_json()


def test_failing_binary_version_smoke_is_typed_unavailable(tmp_path: Path) -> None:
    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(
            binary_probe=lambda: "C:/local/tesseract.exe",
            version_runner=lambda _command: subprocess.CompletedProcess(
                args=[], returncode=1, stdout="", stderr="SECRET runtime failure"
            ),
            module_loader=lambda _name: pytest.fail("adapter import requires binary smoke"),
        ),
    )
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parser_unavailable"
    assert parsed.error_code == "ocr_binary_smoke_failed"
    assert "SECRET" not in parsed.model_dump_json()


@pytest.mark.parametrize(
    ("image_to_data", "expected_code"),
    [
        (
            lambda *_args, **_kwargs: (_ for _ in ()).throw(
                _FakeTesseractError("SECRET missing language data")
            ),
            "ocr_tesseract_error",
        ),
        (lambda *_args, **_kwargs: {"text": ["SECRET malformed"]}, "ocr_malformed_output"),
    ],
)
def test_failing_ocr_runtime_returns_typed_private_outcome(
    tmp_path: Path,
    image_to_data: object,
    expected_code: str,
) -> None:
    module = SimpleNamespace(
        Output=SimpleNamespace(DICT="dict"),
        TesseractError=_FakeTesseractError,
        image_to_data=image_to_data,
    )
    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(
            binary_probe=lambda: "C:/local/tesseract.exe",
            version_runner=lambda _command: subprocess.CompletedProcess(
                args=[], returncode=0, stdout="tesseract 5.5.0", stderr=""
            ),
            module_loader=lambda _name: module,
        ),
    )
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parser_unavailable"
    assert parsed.error_code == expected_code
    assert "SECRET" not in repr(parsed)
    assert "SECRET" not in parsed.model_dump_json()


def test_ocr_local_capability_smoke_is_cached_per_adapter(tmp_path: Path) -> None:
    calls: list[str] = []
    module = SimpleNamespace(
        Output=SimpleNamespace(DICT="dict"),
        TesseractError=_FakeTesseractError,
        image_to_data=lambda *_args, **_kwargs: {
            "text": [],
            "conf": [],
            "left": [],
            "top": [],
            "width": [],
            "height": [],
        },
    )
    adapter = TesseractOcrAdapter(
        binary_probe=lambda: calls.append("binary") or "C:/local/tesseract.exe",
        version_runner=lambda _command: (
            calls.append("version")
            or subprocess.CompletedProcess(
                args=[], returncode=0, stdout="tesseract 5.5.0", stderr=""
            )
        ),
        module_loader=lambda _name: calls.append("module") or module,
    )
    service, store = _service(tmp_path, ocr=adapter)
    payload = _image_bytes()
    artifact = _artifact(payload, "scan.png")
    _persist(store, artifact, payload)

    service.parse(artifact, no_network=True)
    service.parse(artifact, no_network=True)

    assert calls == ["binary", "version", "module"]


def test_image_dimension_limit_rejects_before_ocr_without_raw_leak(tmp_path: Path) -> None:
    called = False

    def recognize(_image: Image.Image) -> list[OcrRegion]:
        nonlocal called
        called = True
        return []

    limits = ImageSafetyLimits(max_width=10, max_height=10, max_pixels=100)
    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(recognizer=recognize, image_limits=limits),
    )
    payload = _image_bytes(metadata_text="SECRET-IMAGE-METADATA")
    artifact = _artifact(payload, "oversized.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "damaged"
    assert parsed.error_code == "image_dimensions_exceeded"
    assert called is False
    assert "SECRET-IMAGE-METADATA" not in repr(parsed)
    assert "SECRET-IMAGE-METADATA" not in parsed.model_dump_json()


def test_pillow_decompression_bomb_warning_is_rejected_before_ocr(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    called = False

    def recognize(_image: Image.Image) -> list[OcrRegion]:
        nonlocal called
        called = True
        return []

    monkeypatch.setattr(Image, "MAX_IMAGE_PIXELS", 100)
    limits = ImageSafetyLimits(max_width=100, max_height=100, max_pixels=10_000)
    service, store = _service(
        tmp_path,
        ocr=TesseractOcrAdapter(recognizer=recognize, image_limits=limits),
    )
    payload = _image_bytes(width=12, height=10)
    artifact = _artifact(payload, "bomb.png")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "damaged"
    assert parsed.error_code == "image_decompression_bomb"
    assert called is False


def test_parser_selection_uses_content_not_extension_or_declared_mime(tmp_path: Path) -> None:
    payload = _docx_bytes()
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "misleading.pdf", mime_type="application/pdf")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == "parsed"
    assert parsed.parser_name == "python-docx"
    assert parsed.detected_mime_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@pytest.mark.parametrize(
    ("payload", "expected_status", "expected_code"),
    [
        (b"not a supported document SECRET-RAW-CONTENT", "unsupported", "unsupported_media_type"),
        (b"%PDF-1.7\nbroken SECRET-RAW-CONTENT", "damaged", "damaged_pdf"),
    ],
)
def test_damaged_or_unsupported_documents_return_typed_private_outcomes(
    tmp_path: Path,
    payload: bytes,
    expected_status: str,
    expected_code: str,
) -> None:
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "document.bin")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    assert parsed.status == expected_status
    assert parsed.error_code == expected_code
    assert "SECRET-RAW-CONTENT" not in repr(parsed)
    assert "SECRET-RAW-CONTENT" not in parsed.model_dump_json()


def test_no_network_guard_blocks_real_socket_attempt() -> None:
    with pytest.raises(NoNetworkViolation, match="socket"):
        with NoNetworkGuard():
            socket.create_connection(("example.com", 443), timeout=0.01)


def test_no_network_guard_detects_model_hub_attempt(monkeypatch: pytest.MonkeyPatch) -> None:
    import huggingface_hub

    called = False

    def fake_download(*_args: object, **_kwargs: object) -> None:
        nonlocal called
        called = True

    monkeypatch.setattr(huggingface_hub, "hf_hub_download", fake_download)
    with pytest.raises(NoNetworkViolation, match="model_hub"):
        with NoNetworkGuard():
            huggingface_hub.hf_hub_download("org/model", "config.json")
    assert called is False


def test_no_network_guard_blocks_model_hub_imported_lazily_and_restores_state(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    package_root = tmp_path / "lazy-module"
    package = package_root / "huggingface_hub"
    package.mkdir(parents=True)
    (package / "__init__.py").write_text(
        "CALLED = False\n"
        "def hf_hub_download(*args, **kwargs):\n"
        "    global CALLED\n"
        "    CALLED = True\n"
        "    return 'network'\n",
        encoding="utf-8",
    )
    monkeypatch.syspath_prepend(str(package_root))
    monkeypatch.delitem(sys.modules, "huggingface_hub", raising=False)
    prior_meta_path = list(sys.meta_path)
    prior_offline = os.environ.get("HF_HUB_OFFLINE")

    with pytest.raises(NoNetworkViolation, match="model_hub"):
        with NoNetworkGuard():
            hub = importlib.import_module("huggingface_hub")
            hub.hf_hub_download("org/model", "config.json")

    assert "huggingface_hub" not in sys.modules
    assert sys.meta_path == prior_meta_path
    assert os.environ.get("HF_HUB_OFFLINE") == prior_offline


def test_no_network_guard_restores_nested_state_after_exception() -> None:
    prior_meta_path = list(sys.meta_path)
    prior_offline = os.environ.get("HF_HUB_OFFLINE")

    with pytest.raises(RuntimeError, match="boom"):
        with NoNetworkGuard():
            with NoNetworkGuard():
                assert os.environ["HF_HUB_OFFLINE"] == "1"
                raise RuntimeError("boom")

    assert sys.meta_path == prior_meta_path
    assert os.environ.get("HF_HUB_OFFLINE") == prior_offline


@pytest.mark.parametrize("inner_raises", [False, True])
def test_nested_no_network_guard_keeps_lazy_hub_reference_blocked_until_outer_exit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    inner_raises: bool,
) -> None:
    package_root = tmp_path / "nested-lazy-module"
    package = package_root / "huggingface_hub"
    package.mkdir(parents=True)
    (package / "__init__.py").write_text(
        "CALLED = False\n"
        "def hf_hub_download(*args, **kwargs):\n"
        "    global CALLED\n"
        "    CALLED = True\n"
        "    return 'network'\n",
        encoding="utf-8",
    )
    monkeypatch.syspath_prepend(str(package_root))
    monkeypatch.delitem(sys.modules, "huggingface_hub", raising=False)

    with NoNetworkGuard():
        try:
            with NoNetworkGuard():
                hub = importlib.import_module("huggingface_hub")
                saved_download = hub.hf_hub_download
                if inner_raises:
                    raise RuntimeError("inner exit")
        except RuntimeError as exc:
            assert inner_raises is True
            assert str(exc) == "inner exit"

        with pytest.raises(NoNetworkViolation, match="model_hub"):
            hub.hf_hub_download("org/model", "config.json")
        with pytest.raises(NoNetworkViolation, match="model_hub"):
            saved_download("org/model", "config.json")
        assert hub.CALLED is False

    assert hub.hf_hub_download("org/model", "config.json") == "network"
    assert hub.CALLED is True


def test_docling_registration_is_lazy_and_requires_dependency_cache_and_offline_smoke(
    tmp_path: Path,
) -> None:
    cache = tmp_path / "docling-cache"
    cache.mkdir()
    (cache / "weights.bin").write_bytes(b"local model")
    (cache / "model-manifest.json").write_text(
        json.dumps(
            {
                "files": [
                    {
                        "path": "weights.bin",
                        "size": len(b"local model"),
                        "sha256": sha256(b"local model").hexdigest(),
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    calls: list[str] = []

    parser = DoclingDocumentParser.try_create(
        model_cache=cache,
        dependency_probe=lambda name: calls.append(name) or True,
        offline_smoke=lambda _cache: calls.append("smoke") or True,
    )

    assert parser is not None
    assert calls == ["docling", "smoke"]
    assert DoclingDocumentParser.try_create(
        model_cache=cache,
        dependency_probe=lambda _name: False,
        offline_smoke=lambda _cache: pytest.fail("smoke must stay lazy"),
    ) is None
    assert DoclingDocumentParser.try_create(
        model_cache=tmp_path / "missing",
        dependency_probe=lambda _name: True,
        offline_smoke=lambda _cache: pytest.fail("smoke requires cache"),
    ) is None
    assert DoclingDocumentParser.try_create(
        model_cache=cache,
        dependency_probe=lambda _name: True,
        offline_smoke=lambda _cache: False,
    ) is None


def test_docling_smoke_is_forced_offline_and_tampered_cache_is_rejected(
    tmp_path: Path,
) -> None:
    payload = b"recorded local model"
    cache = tmp_path / "docling-cache"
    cache.mkdir()
    (cache / "weights.bin").write_bytes(payload)
    (cache / "model-manifest.json").write_text(
        json.dumps(
            {
                "files": [
                    {
                        "path": "weights.bin",
                        "size": len(payload),
                        "sha256": sha256(payload).hexdigest(),
                    }
                ]
            }
        ),
        encoding="utf-8",
    )

    parser = DoclingDocumentParser.try_create(
        model_cache=cache,
        dependency_probe=lambda _name: True,
        offline_smoke=lambda _cache: os.environ.get("HF_HUB_OFFLINE") == "1",
    )
    assert parser is not None

    (cache / "weights.bin").write_bytes(b"tampered")
    assert DoclingDocumentParser.try_create(
        model_cache=cache,
        dependency_probe=lambda _name: True,
        offline_smoke=lambda _cache: pytest.fail("tampered cache must fail before smoke"),
    ) is None


def test_text_models_contain_only_references_hashes_counts_and_confidence(tmp_path: Path) -> None:
    secret = "ACME SECRET PIPELINE 123456"
    payload = _pdf_with_text(secret)
    service, store = _service(tmp_path)
    artifact = _artifact(payload, "deck.pdf")
    _persist(store, artifact, payload)

    parsed = service.parse(artifact, no_network=True)

    dumped = parsed.model_dump_json()
    assert secret not in dumped
    assert secret not in repr(parsed)
    block = parsed.text_blocks[0]
    assert block.text_ref == block.content_hash
    assert block.char_count == len(secret)
    assert store.read_bytes(block.text_ref).decode() == secret


def _service(
    tmp_path: Path,
    *,
    ocr: TesseractOcrAdapter | None = None,
) -> tuple[StartupParsingService, LocalArtifactStore]:
    store = LocalArtifactStore(tmp_path / "artifact-store")
    return StartupParsingService(artifact_store=store, ocr_parser=ocr), store


def _artifact(payload: bytes, source: str, mime_type: str = "application/octet-stream") -> Artifact:
    from hashlib import sha256

    digest = sha256(payload).hexdigest()
    return Artifact(
        id=uuid4(),
        case_id=UUID("aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa"),
        content_hash=digest,
        mime_type=mime_type,
        source=source,
        retrieved_at=datetime.now(UTC),
        source_snapshot_hash=digest,
        sensitivity=SensitivityClass.RESTRICTED,
    )


def _persist(store: LocalArtifactStore, artifact: Artifact, payload: bytes) -> None:
    store.put_bytes(
        payload,
        media_type=artifact.mime_type,
        artifact_id=artifact.id,
        source_snapshot_hash=artifact.source_snapshot_hash,
        sensitivity=artifact.sensitivity,
    )


def _pdf_with_text(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page(width=300, height=200)
    page.insert_text((36, 72), text)
    payload = document.tobytes()
    document.close()
    return payload


def _pdf_with_table() -> bytes:
    document = fitz.open()
    page = document.new_page(width=300, height=200)
    for x in (30, 150, 250):
        page.draw_line((x, 30), (x, 110))
    for y in (30, 70, 110):
        page.draw_line((30, y), (250, y))
    page.insert_text((40, 55), "Metric")
    page.insert_text((160, 55), "Value")
    page.insert_text((40, 95), "ARR")
    page.insert_text((160, 95), "1200")
    payload = document.tobytes()
    document.close()
    return payload


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Board composition")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Director"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Ada"
    table.cell(1, 1).text = "Chair"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class _FakeTesseractError(RuntimeError):
    pass


def _image_bytes(
    *,
    mode: str = "RGB",
    width: int = 12,
    height: int = 10,
    metadata_text: str | None = None,
) -> bytes:
    image = Image.new(mode, (width, height))
    output = BytesIO()
    pnginfo = None
    if metadata_text is not None:
        pnginfo = PngImagePlugin.PngInfo()
        pnginfo.add_text("comment", metadata_text)
    image.save(output, format="PNG", pnginfo=pnginfo)
    return output.getvalue()
